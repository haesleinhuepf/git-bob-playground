#!/usr/bin/env python3
"""
Randomize the positions of table cell texts in a DOCX bingo sheet and save multiple versions.

Usage:
  python randomize_bingo.py --input "Prompt Engineering Bingo.docx" --n 20 --outdir out --seed 42
"""

from __future__ import annotations
import argparse
import os
import random
from typing import List, Optional

try:
    from docx import Document
    from docx.table import _Cell, Table
except Exception as e:
    raise RuntimeError(
        "python-docx is required. Install with `pip install python-docx`."
    ) from e


def table_cell_texts(table: Table) -> List[str]:
    """Collect plain text from all cells of a table.

    Parameters
    ----------
    table : docx.table.Table
        The table whose cell texts are collected.

    Returns
    -------
    list of str
        Text content of each cell in row-major order.
    """
    return [cell.text for row in table.rows for cell in row.cells]


def assign_texts_to_table(table: Table, texts: List[str]) -> None:
    """Assign a list of texts back to a table's cells in row-major order.

    Parameters
    ----------
    table : docx.table.Table
        The table to write texts into.
    texts : list of str
        Texts to assign. Must match number of cells.
    """
    cells = [cell for row in table.rows for cell in row.cells]
    if len(cells) != len(texts):
        raise ValueError(f"Number of texts ({len(texts)}) does not match number of cells ({len(cells)}).")
    for cell, txt in zip(cells, texts):
        cell.text = txt or ""


def index_of_largest_table(doc: Document) -> int:
    """Find the index of the largest table by number of cells.

    Parameters
    ----------
    doc : docx.document.Document
        The document to inspect.

    Returns
    -------
    int
        Index of the largest table. Returns 0 if there are no tables.
    """
    if not doc.tables:
        return 0
    sizes = [len([c for r in t.rows for c in r.cells]) for t in doc.tables]
    return max(range(len(sizes)), key=lambda i: sizes[i])


def randomize_tables_in_document(doc: Document, rng: random.Random, all_tables: bool = False) -> None:
    """Randomize cell texts within tables of a document.

    Parameters
    ----------
    doc : docx.document.Document
        Document to modify in-place.
    rng : random.Random
        Random number generator to use for shuffling.
    all_tables : bool, default False
        If True, randomize all tables. If False, randomize only the largest table.
    """
    if not doc.tables:
        return
    table_indices = range(len(doc.tables)) if all_tables else [index_of_largest_table(doc)]
    for idx in table_indices:
        table = doc.tables[idx]
        texts = table_cell_texts(table)
        shuffled = texts[:]  # copy
        rng.shuffle(shuffled)
        assign_texts_to_table(table, shuffled)


def save_randomized_versions(input_path: str, outdir: str, n: int = 20, seed: Optional[int] = None, all_tables: bool = False) -> List[str]:
    """Create N randomized versions of a DOCX by shuffling table cell texts.

    Parameters
    ----------
    input_path : str
        Path to the input DOCX file (bingo template).
    outdir : str
        Directory to save randomized files into. Will be created if missing.
    n : int, default 20
        Number of randomized documents to create.
    seed : int, optional
        Random seed for reproducibility. If None, system randomness is used.
    all_tables : bool, default False
        If True, randomize all tables; otherwise only the largest table.

    Returns
    -------
    list of str
        Paths to the saved randomized documents.
    """
    os.makedirs(outdir, exist_ok=True)
    rng = random.Random(seed)
    base = os.path.splitext(os.path.basename(input_path))[0]
    out_paths = []
    for i in range(1, n + 1):
        doc = Document(input_path)
        randomize_tables_in_document(doc, rng, all_tables=all_tables)
        out_path = os.path.join(outdir, f"{base}_randomized_{i:02d}.docx")
        doc.save(out_path)
        out_paths.append(out_path)
    return out_paths


def main():
    parser = argparse.ArgumentParser(description="Randomize DOCX bingo table cell contents and create multiple versions.")
    parser.add_argument("--input", "-i", required=False, default="test.docx", help="Path to the input DOCX (default: test.docx).")
    parser.add_argument("--outdir", "-o", required=False, default="randomized", help="Output directory (default: ./randomized).")
    parser.add_argument("--n", "-n", type=int, required=False, default=20, help="Number of randomized versions to create (default: 20).")
    parser.add_argument("--seed", type=int, required=False, default=None, help="Random seed for reproducibility (default: None).")
    parser.add_argument("--all-tables", action="store_true", help="Randomize all tables (default: only the largest table).")
    args = parser.parse_args()

    paths = save_randomized_versions(args.input, args.outdir, n=args.n, seed=args.seed, all_tables=args.all_tables)
    print(f"Saved {len(paths)} files:")
    for p in paths:
        print(f" - {p}")


if __name__ == "__main__":
    main()
